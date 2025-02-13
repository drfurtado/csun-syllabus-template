from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Set page margins
sections = doc.sections
for section in sections:
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Create header with logo and text
    header = section.header
    # Add a 2-column table in the header with full width
    table = header.add_table(rows=1, cols=2, width=Twips(section.page_width - section.left_margin - section.right_margin))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True
    
    # Left cell: Logo
    left_cell = table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_run = left_para.add_run()
    # Make logo height match two lines of 12pt text plus spacing
    left_run.add_picture('csun-logo-syllabi.png', height=Pt(36))  # 12pt * 2 lines + spacing
    
    # Right cell: Course info
    right_cell = table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    course_info = right_para.add_run('[Course Number]\n[Term]')
    course_info.font.name = 'Arial'
    course_info.font.size = Pt(12)
    right_para.paragraph_format.space_after = Pt(0)
    right_para.paragraph_format.space_before = Pt(0)
    
    # Make header table borderless
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                for border in ['top', 'right', 'bottom', 'left']:
                    setattr(cell._tc.get_or_add_tcPr(), 'w:'+border, None)
    
    # Add space after header
    spacing_para = header.add_paragraph()
    spacing_para.space_after = Pt(24)

# Define CSUN styles
def set_csun_heading_style(style, size):
    font = style.font
    font.name = 'Arial'
    font.size = Pt(size)
    font.bold = True
    font.color.rgb = RGBColor(0xD2, 0x20, 0x30)  # CSUN Red

def set_csun_normal_style(style):
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5

# Configure List Styles for single spacing
list_style = doc.styles['List Paragraph']
list_style.font.name = 'Arial'
list_style.font.size = Pt(12)
list_style.paragraph_format.line_spacing = 1.0
list_style.paragraph_format.left_indent = Inches(0.25)

bullet_style = doc.styles['List Bullet']
bullet_style.base_style = list_style
bullet_style.font.name = 'Arial'
bullet_style.font.size = Pt(12)
bullet_style.paragraph_format.line_spacing = 1.0
bullet_style.paragraph_format.left_indent = Inches(0.25)

# Apply styles
heading1_style = doc.styles['Heading 1']
set_csun_heading_style(heading1_style, 16)

heading2_style = doc.styles['Heading 2']
set_csun_heading_style(heading2_style, 14)

normal_style = doc.styles['Normal']
set_csun_normal_style(normal_style)

# Save the reference document
doc.save('csun-syllabus-reference.docx')